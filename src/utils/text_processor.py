import re
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
# import markdown2
# from weasyprint import HTML

def clean_llm_output(text: str) -> str:
    """
    Clean LLM output by removing thinking/reasoning content.
    
    Args:
        text (str): The raw LLM output
        
    Returns:
        str: Cleaned text without thinking content
    """
    # Remove content between <think> tags
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)
    
    # Remove any remaining <think> tags
    text = re.sub(r'</?think>', '', text)
    
    # Clean up any extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.strip()
    
    return text

def markdown_to_docx(markdown_text: str) -> bytes:
    """
    Convert markdown text to a formatted DOCX document.
    
    Args:
        markdown_text (str): The markdown text to convert
        
    Returns:
        bytes: The DOCX file as bytes
    """
    # Convert markdown to HTML
    html = markdown.markdown(markdown_text)
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Research Paper', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    for line in markdown_text.split('\n'):
        if line.startswith('#'):
            # Handle headings
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=level)
        else:
            # Handle regular paragraphs
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

# def markdown_to_pdf(markdown_text: str) -> bytes:
#     """
#     Convert markdown text to a formatted PDF document.
    
#     Args:
#         markdown_text (str): The markdown text to convert
        
#     Returns:
#         bytes: The PDF file as bytes
#     """
#     # Convert markdown to HTML
#     html_content = markdown2.markdown(markdown_text)

#     # Convert HTML to PDF
#     pdf_bytes = io.BytesIO()
#     HTML(string=html_content).write_pdf(pdf_bytes)
#     pdf_bytes.seek(0)

#     return pdf_bytes.getvalue()